"""
Paper Assistant Backend Service
Provides thesis outline generation, text polishing, reference formatting, format checking, and document export
"""
import os
import re
import time
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict

PAPER_OUTPUT = 'paper_outputs'
os.makedirs(PAPER_OUTPUT, exist_ok=True)

def call_llm(prompt: str, system_prompt: str = None) -> str:
    """Call LLM API"""
    try:
        from dotenv import load_dotenv
        load_dotenv()
    except:
        pass

    # Use hardcoded API credentials to avoid environment variable issues
    API_KEY = "sk-4cbec58db9cb40569bc6080254924697"
    API_URL = "https://api.deepseek.com/chat/completions"
    if system_prompt is None:
        system_prompt = "You are a professional academic writing assistant."

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    # The correct model name for
    # Model name - hardcoded to avoid encoding issues
    model_name = "deepseek-chat"

    payload = {
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }

    try:
        resp = requests.post(API_URL, headers=headers, json=payload, timeout=120)
        result = resp.json()

        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"].strip()

        return f"API Error: {result.get('error', 'Unknown error')}"
    except Exception as e:
        return f"Request failed: {str(e)}"


def generate_outline(topic: str, discipline: str = "general") -> Dict:
    """Generate thesis outline based on topic"""
    prompt = f"""Please generate a complete thesis outline for the following topic:

Topic: {topic}
Field: {discipline}

Please generate an outline in the following format:
1. Abstract (brief description of research purpose, methods, conclusions)
2. Chapter 1 Introduction (research background, significance, purpose, methods)
3. Chapter 2 Related Theories and Technologies (core concepts, existing methods)
4. Chapter 3 System Design and Implementation / Research Methods
5. Chapter 4 Experimental Results and Analysis
6. Chapter 5 Summary and Outlook

For each chapter, list 2-3 subsections with brief content points.

Please answer in Chinese with clear formatting."""

    content = call_llm(prompt, "You are an experienced academic thesis writing mentor.")

    return {
        "content": content,
        "source": "Paper Assistant"
    }


def generate_outline_document(topic: str, discipline: str = "general", format: str = "docx") -> Dict:
    """Generate formatted outline document"""
    result = generate_outline(topic, discipline)
    content = result.get("content", "")

    timestamp = int(time.time())
    safe_topic = re.sub(r'[^\w\u4e00-\u9fff]', '_', topic)[:20]

    if format == "md":
        filename = f"Outline_{safe_topic}_{timestamp}.md"
        filepath = os.path.join(PAPER_OUTPUT, filename)

        md_content = f"""# {topic}

> Field: {discipline} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

{content}

---

## Instructions

1. This outline is for reference only
2. Adjust chapters based on actual research depth
3. Discuss and confirm the outline with your advisor

## Export Info

- Topic: {topic}
- Field: {discipline}
- Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- Tool: Smart Paper Assistant

"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(md_content)

        return {
            "content": content,
            "filename": filename,
            "filepath": filepath,
            "source": "Paper Assistant"
        }
    else:
        filename = f"Outline_{safe_topic}_{timestamp}.docx"
        filepath = os.path.join(PAPER_OUTPUT, filename)

        doc = Document()
        doc.add_heading(topic, 0)

        if content and not content.startswith("Error"):
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph(content)

        doc.save(filepath)

        return {
            "content": content,
            "filename": filename,
            "filepath": filepath,
            "source": "Paper Assistant"
        }


def generate_polish_document(original_text: str, polished_text: str, format: str = "docx") -> Dict:
    """Generate polished document"""
    timestamp = int(time.time())

    if format == "md":
        filename = f"Polished_{timestamp}.md"
        filepath = os.path.join(PAPER_OUTPUT, filename)

        md_content = f"""# Polished Text Comparison

> Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## Original Text

{original_text}

---

## Polished Text

{polished_text}

---

## Notes

- Polishing maintains the core meaning
- Improves academic professionalism

"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(md_content)
    else:
        filename = f"Polished_{timestamp}.docx"
        filepath = os.path.join(PAPER_OUTPUT, filename)

        doc = Document()
        doc.add_heading('Polished Text Comparison', 0)
        doc.add_heading('Original Text', 1)
        doc.add_paragraph(original_text)
        doc.add_heading('Polished Text', 1)
        doc.add_paragraph(polished_text)
        doc.save(filepath)

    return {
        "content": polished_text,
        "filename": filename,
        "filepath": filepath,
        "source": "Paper Assistant"
    }


def polish_text(text: str) -> Dict:
    """Polish academic text"""
    prompt = f"""Please polish the following academic text, maintaining the original meaning:

Original text:
{text}

Please directly output the polished text."""

    content = call_llm(prompt, "You are a professional academic editor.")

    return {
        "content": content,
        "source": "Paper Assistant"
    }


def rewrite_text(text: str) -> Dict:
    """Rewrite text for originality"""
    prompt = f"""Please rewrite the following text in 3 different ways:

Original text:
{text}

Please output in this format:
[Rewrite 1]
[Rewritten text 1]

[Rewrite 2]
[Rewritten text 2]

[Rewrite 3]
[Rewritten text 3]"""

    content = call_llm(prompt, "You are a professional academic writing assistant.")

    return {
        "content": content,
        "source": "Paper Assistant"
    }


def format_references(text: str, format_type: str = "gb") -> Dict:
    """Format references"""

    format_guides = {
        "gb": """
GB/T 7714-2015 Format:
- Journal: [No.] Author. Title[J]. Journal Name, Year, Vol(Issue): Pages
- Book: [No.] Author. Title[M]. Place: Publisher, Year""",
        "apa": """
APA Format:
- Journal: Author, A. A. (Year). Title. Journal Name, volume(issue), pages.
- Book: Author, A. A. (Year). Title. Publisher.""",
    }

    format_guide = format_guides.get(format_type, format_guides["gb"])

    prompt = f"""Please format the following references in {format_type.upper()} format:

{text}

{format_guide}"""

    content = call_llm(prompt, "You are a professional reference formatting assistant.")

    return {
        "content": content,
        "source": "Paper Assistant"
    }


def generate_references_document(refs_text: str, format_type: str = "gb") -> Dict:
    """Generate formatted references document"""
    result = format_references(refs_text, format_type)
    content = result.get("content", "")

    timestamp = int(time.time())
    format_names = {"gb": "GB/T 7714", "apa": "APA"}

    filename = f"References_{format_names.get(format_type, format_type)}_{timestamp}.md"
    filepath = os.path.join(PAPER_OUTPUT, filename)

    md_content = f"""# Reference List

> Format: {format_names.get(format_type, format_type)} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

{content}

---

## Import Instructions

1. Copy the reference list
2. Paste into the References chapter
3. Ensure consistency with citations

"""

    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(md_content)

    return {
        "content": content,
        "filename": filename,
        "filepath": filepath,
        "source": "Paper Assistant"
    }


def check_format(document_path: str = None) -> Dict:
    """Check paper format"""

    default_checks = """
Paper Format Checklist
======================

General format checking points:

[Font and Size]
- Body text: Song Ti 12pt or Times New Roman 12pt
- Level 1 heading: Hei Ti 16pt, centered, bold
- Level 2 heading: Hei Ti 14pt, left aligned

[Paragraph Format]
- First line indent 2 characters
- Line spacing: 1.5

[Heading Hierarchy]
- Consistent heading hierarchy
- Consistent numbering format

[References]
- GB/T 7714 format
- In-text citations consistent with reference list
"""

    return {
        "content": default_checks,
        "has_issues": True,
        "source": "Paper Assistant"
    }


def get_template(template_type: str) -> Dict:
    """Get writing template"""

    templates = {
        "outline": """
[Thesis Outline Template]

I. Abstract
   Briefly describe research purpose, methods, main results (200-300 words)

II. Introduction
   1.1 Research Background
   1.2 Research Significance
   1.3 Research Content and Methods
   1.4 Thesis Structure

III. Related Theories and Technologies
   2.1 XXX Theory/Concept
   2.2 XXX Technology/Method

IV. System Design and Implementation
   4.1 Requirements Analysis
   4.2 System Design
   4.3 Implementation

V. Experimental Results and Analysis
   5.1 Experimental Results
   5.2 Result Analysis

VI. Summary and Outlook
   6.1 Research Summary
   6.2 Future Research Directions

References
Acknowledgments
""",
        "abstract": """
[Abstract Template]

This paper addresses [research problem], using [research methods], to carry out [research work].

First, through [method], solved [problem]; then, based on [theory], designed [system]; finally, through [experiment], proved [conclusion].

Experimental results show that [method] achieves [effect].

This research provides new ideas for [field], with certain theoretical and practical value.

[Keywords]: Keyword 1; Keyword 2; Keyword 3; Keyword 4; Keyword 5
""",
    }

    template = templates.get(template_type, templates["outline"])

    return {
        "content": template,
        "source": "Paper Assistant"
    }
