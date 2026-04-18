import os
import requests
import json
import re
import time
import urllib.parse # 🚨 新增：用于 URL 转码
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
from commute_helper import is_commute_request, commute_helper
from datasheet_rag import datasheet_rag
from paper_helper import (
    generate_outline, polish_text, rewrite_text,
    format_references, check_format, get_template,
    generate_outline_document, generate_polish_document,
    generate_references_document
)

from docx import Document
from docx.shared import Pt, RGBColor

# 🚨 重要修复：手动加载环境变量（避免 dotenv 编码问题）
try:
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    if os.path.exists(env_path):
        with open(env_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    os.environ[key.strip()] = value.strip()
        print(f"[MAIN] 环境变量已手动加载")
    else:
        print(f"[MAIN] .env 文件不存在: {env_path}")
        load_dotenv(override=True)
except Exception as e:
    print(f"[MAIN] 环境变量加载失败: {e}")
    load_dotenv(override=True)

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def get_model_config():
    return {
        "deepseek": {"url": os.getenv("DEEPSEEK_URL"), "key": os.getenv("DEEPSEEK_KEY"), "model": "deepseek-chat"},
        "kimi": {"url": os.getenv("KIMI_URL"), "key": os.getenv("KIMI_KEY"), "model": "moonshot-v1-8k"},
        "doubao": {"url": "https://ark.cn-beijing.volces.com/api/v3/responses", "key": os.getenv("DOUBAO_KEY"), "model": "doubao-seed-2-0-pro-260215"}
    }

def process_document_with_agent(filepath, original_filename, question, system_prompt, config):
    try:
        doc = Document(filepath)
        paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs_text)

        agent_prompt = f"""你是一个专业的文档排版与润色智能体。
用户指令：{question}

文档内容如下：
{full_text}

请严格按以下JSON格式返回修改建议（不要返回任何Markdown解释，只返回纯JSON字符串）：
{{
  "rewrite_tasks": [
    {{"target_text": "需要被替换的原句子", "new_text": "润色后的新句子"}}
  ],
  "format_tasks": [
    {{"target": "body_text", "font_name": "SimSun", "font_size": 14, "bold": false, "italic": false, "underline": false, "color": "black"}},
    {{"target": "text:特定文本", "font_size": 16, "bold": true, "color": "red"}}
  ]
}}

📋 format_tasks 字段说明：
- target: 修改对象
  * "body_text" = 修改整个文档的所有文本
  * "text:具体文本内容" = 只修改特定的文本（例如 "text:第一段" 只修改"第一段"这个文本）
- font_name: 字体名称（必填）
  * SimSun = 宋体
  * SimHei = 黑体
  * Arial = 英文Arial字体
  * Courier New = 等宽字体
  * Comic Sans = 手写体
- font_size: 数字，字体大小（如 12、14、16、20）
- bold: true=加粗，false=不加粗
- italic: true=斜体，false=不斜体
- underline: true=下划线，false=无下划线
- color: 字体颜色，可用值包括：
  * black（黑色）、red（红色）、blue（蓝色）、green（绿色）
  * yellow（黄色）、orange（橙色）、purple（紫色）、brown（棕色）
  * gray/grey（灰色）、white（白色）、pink（粉色）、cyan（青色）
  * magenta（品红）、lime（青绿）、navy（深蓝）、teal（深青）、olive（橄榄绿）

📌 重要指南：
1. 用户要求「把第二段改成红色加粗」时，你应该返回：
   {{"target": "text:第二段", "color": "red", "bold": true, "font_name": "SimHei", "font_size": 14}}

2. 用户要求「把整个文档改成黑体12号」时，返回：
   {{"target": "body_text", "font_name": "SimHei", "font_size": 12}}

3. 如果图片中的用户指令没有提及字体，保持原样（不填相关字段）
4. 如果用户没有要求修改格式，format_tasks 可以为空数组 []
5. 重点：所有布尔值（bold、italic、underline）必须是 true 或 false，不能是 1/0 或 yes/no
"""

        payload = {
            "model": config["model"],
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": agent_prompt}]
        }

        headers = {"Authorization": f"Bearer {config['key']}", "Content-Type": "application/json"}
        resp = requests.post(config["url"], headers=headers, json=payload, timeout=300)
        result = resp.json()
        
        if "choices" not in result or len(result["choices"]) == 0:
            err_msg = result.get("error", {}).get("message", "API返回的数据异常")
            return jsonify({"content": f"大模型接口拒绝了请求，原因：\n{err_msg}", "source": "系统"})

        answer = result["choices"][0]["message"].get("content", "").strip()
        start_idx = answer.find('{')
        end_idx = answer.rfind('}')
        
        if start_idx != -1 and end_idx != -1:
            json_str = answer[start_idx:end_idx+1]
            json_str = re.sub(r',\s*([\]}])', r'\1', json_str)
            instructions = json.loads(json_str, strict=False)
        else:
            raise ValueError(f"大模型没有返回包含 {{}} 的有效 JSON 数据。原回复：\n{answer}")

        for task in instructions.get("rewrite_tasks", []):
            old_text = task.get("target_text", "")
            new_text = task.get("new_text", "")
            if old_text and new_text:
                for p in doc.paragraphs:
                    if old_text in p.text:
                        # 🚨 核心修复：使用 replace_text 确保格式正确应用
                        for run in p.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

        for task in instructions.get("format_tasks", []):
            target = task.get("target")
            font_name = task.get("font_name")
            font_size = task.get("font_size")
            bold = task.get("bold", False)
            italic = task.get("italic", False)
            underline = task.get("underline", False)
            color = task.get("color", "black")
            
            # 🚨 颜色映射表：支持常见的中文和英文颜色名称
            color_map = {
                "black": (0, 0, 0),
                "red": (255, 0, 0),
                "blue": (0, 0, 255),
                "green": (0, 128, 0),
                "yellow": (255, 255, 0),
                "orange": (255, 165, 0),
                "purple": (128, 0, 128),
                "brown": (165, 42, 42),
                "gray": (128, 128, 128),
                "grey": (128, 128, 128),
                "white": (255, 255, 255),
                "pink": (255, 192, 203),
                "cyan": (0, 255, 255),
                "magenta": (255, 0, 255),
                "lime": (0, 255, 0),
                "navy": (0, 0, 128),
                "teal": (0, 128, 128),
                "olive": (128, 128, 0),
            }
            
            print(f"📝 应用格式: target={target}, font={font_name}, size={font_size}, bold={bold}, color={color}")
            
            if target == "body_text":
                # 遍历所有段落和run，应用格式
                for p in doc.paragraphs:
                    for run in p.runs:
                        # 应用字体名称
                        if font_name:
                            run.font.name = font_name
                            run.font.ascii_theme_color = None  # 清除主题颜色以使用自定义颜色
                        
                        # 应用字体大小
                        if font_size:
                            run.font.size = Pt(font_size)
                        
                        # 应用加粗
                        if bold:
                            run.font.bold = True
                        
                        # 应用斜体
                        if italic:
                            run.font.italic = True
                        
                        # 应用下划线
                        if underline:
                            run.font.underline = True
                        
                        # 应用颜色（这是关键！）
                        if color and color.lower() in color_map:
                            rgb = color_map[color.lower()]
                            run.font.color.rgb = RGBColor(*rgb)
                            print(f"✅ 颜色已应用: {color} = RGB{rgb}")
            
            # 支持特定文本的格式修改
            elif target and target.startswith("text:"):
                target_text = target.replace("text:", "").strip()
                print(f"🔍 寻找特定文本: {target_text}")
                
                for p in doc.paragraphs:
                    for run in p.runs:
                        if target_text in run.text:
                            if font_name:
                                run.font.name = font_name
                            if font_size:
                                run.font.size = Pt(font_size)
                            if bold:
                                run.font.bold = True
                            if italic:
                                run.font.italic = True
                            if underline:
                                run.font.underline = True
                            if color and color.lower() in color_map:
                                rgb = color_map[color.lower()]
                                run.font.color.rgb = RGBColor(*rgb)
                            print(f"✅ 文本格式已应用: {target_text}")

        print("🎉 文档格式处理完成！")

        timestamp = int(time.time())
        output_filename = f"修改后_{timestamp}_{original_filename}"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        doc.save(output_path)

        # 🚨 核心修复：用固定 IP 而非 request.host，避免手机端拿到无效地址
        server_ip = os.getenv("SERVER_IP", "192.168.139.221")
        server_port = os.getenv("FLASK_PORT", "5000")
        encoded_filename = urllib.parse.quote(output_filename)
        file_url = f"http://{server_ip}:{server_port}/api/download/{encoded_filename}"

        return jsonify({
            "content": "✅ 文档已智能排版与润色完毕！\n请点击下方文件卡片导出至手机。",
            "source": "文档专家",
            "file_path": file_url  
        })

    except json.JSONDecodeError:
        return jsonify({"content": f"处理失败：提取出的指令不是标准JSON格式。\n大模型原回复：\n{answer}", "source": "系统"})
    except Exception as e:
        return jsonify({"content": f"文档处理本地执行异常：{str(e)}", "source": "系统"})

@app.route('/api/chat', methods=['POST'])
def chat():
    # 意图识别阶段跳过插件路由，避免污染识别结果
    skip_plugin_routing = request.args.get('skip_plugin') == '1'
    if request.content_type and 'multipart/form-data' in request.content_type:
        data = request.form
        question = data.get('question', '').strip()
        system_prompt = data.get('system_prompt', '你是一个AI助手')
        model = data.get('model', 'kimi') 
        
        document = request.files.get('document')
        if not document or document.filename == '':
            return jsonify({"content": "未接收到有效的文档文件", "source": "错误"})

        try:
            original_filename = document.filename.replace('/', '_').replace('\\', '_') 
            safe_input_filename = f"upload_{int(time.time())}_{original_filename}"
            input_path = os.path.join(UPLOAD_FOLDER, safe_input_filename)
            document.save(input_path)
            
            config = get_model_config()[model]
            return process_document_with_agent(input_path, original_filename, question, system_prompt, config)
        except Exception as e:
            return jsonify({"content": f"文件保存或处理失败: {str(e)}", "source": "错误"})

    else:
        data = request.json
        question = data.get('question', '').strip()
        image_data = data.get('image_data')
        model = data.get('model', 'deepseek')
        system_prompt = data.get('system_prompt', '你叫小陆，是圆圆的贴心助手。')

        # Commute assistant routing（意图识别阶段跳过，避免误判）
        if not skip_plugin_routing and is_commute_request(question):
            return jsonify(commute_helper.handle_request(question))

        # ============ 论文/数据手册自动导出（意图识别阶段也执行）============
        # 论文助手意图识别
        paper_keywords = ['论文', '大纲', '开题报告', '毕业设计', '毕业论文']
        paper_export_keywords = ['导出', '输出', '生成文档', '生成word', '生成md', '下载', 'word格式', 'markdown格式', 'md格式', 'docx', '文档']
        is_paper_request = any(kw in question for kw in paper_keywords)
        is_paper_export = any(kw in question.lower() for kw in paper_export_keywords)

        if is_paper_request:
            print(f"[PAPER] 论文请求: topic='{question}'")
            # 提取主题（去掉关键词）
            topic = question
            for kw in ['论文', '大纲', '开题报告', '毕业设计', '毕业论文', '导出', '输出', '生成文档', '生成word', '生成md', '下载', 'word格式', 'markdown格式', 'md格式', 'docx']:
                topic = topic.replace(kw, '').replace(kw.upper(), '').replace(kw.capitalize(), '')
            topic = topic.strip()

            fmt = 'md' if 'md' in question.lower() or 'markdown' in question.lower() else 'docx'
            server_ip = os.getenv("SERVER_IP", "192.168.2.6")
            server_port = os.getenv("FLASK_PORT", "5000")

            # 1. 先调用 AI 生成大纲（论文助手核心）
            ai_answer = ""
            if topic:
                ai_result = generate_outline(topic, 'general')
                ai_answer = ai_result.get("content", "")
                print(f"[PAPER] AI生成完成，字符数: {len(ai_answer)}")

            # 2. 如果需要导出文档，生成文档（Word/Markdown）
            file_url = None
            if is_paper_export and topic:
                doc_result = generate_outline_document(topic, 'general', fmt)
                filepath = doc_result.get("filepath", "")
                if filepath and os.path.exists(filepath):
                    filename = os.path.basename(filepath)
                    file_url = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
                    print(f"[PAPER] 文档已生成: {filepath}")

            # 3. 同时返回消息内容和文件下载链接
            resp_content = ""
            if ai_answer and not ai_answer.startswith("错误：") and not ai_answer.startswith("请求失败"):
                resp_content = ai_answer
            elif topic:
                resp_content = f"正在为您生成关于「{topic}」的论文大纲...\n\n（文档生成中，请稍候）"

            if file_url:
                fmt_name = "Word文档(.docx)" if fmt == 'docx' else "Markdown(.md)"
                resp_content += f"\n\n📄 {fmt_name}已生成！点击下方卡片下载文件："

            return jsonify({
                "content": resp_content,
                "source": "论文助手",
                "file_url": file_url,
                "file_path": file_url,
            })

        # 数据手册助手意图识别（嵌入式/芯片相关）
        datasheet_keywords = ['数据手册', 'datasheet', '芯片', '单片机', 'stm32', 'esp32', 'gpio', '寄存器', '嵌入式']
        datasheet_export_keywords = ['导出', '输出', '生成文档', '生成word', '生成md', '下载', 'word格式', 'markdown格式', 'md格式', 'docx', '文档']
        is_datasheet_request = any(kw in question.lower() for kw in datasheet_keywords)
        is_datasheet_export = any(kw in question.lower() for kw in datasheet_export_keywords)

        if is_datasheet_request:
            print(f"[DATASHEET] 数据手册请求: question='{question}'")
            fmt = 'md' if 'md' in question.lower() or 'markdown' in question.lower() else 'docx'
            server_ip = os.getenv("SERVER_IP", "192.168.2.6")
            server_port = os.getenv("FLASK_PORT", "5000")

            # 1. 先调用 RAG 回答问题（数据手册核心）
            rag_result = datasheet_rag.ask(question)
            rag_answer = rag_result.get("content", "")
            print(f"[DATASHEET] RAG回答完成，字符数: {len(rag_answer)}")

            # 2. 如果需要导出文档，生成文档
            file_url = None
            if is_datasheet_export:
                doc_result = datasheet_rag.export_to_document(question, rag_answer, fmt)
                filepath = doc_result.get("filepath", "")
                if filepath and os.path.exists(filepath):
                    filename = os.path.basename(filepath)
                    file_url = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
                    print(f"[DATASHEET] 文档已生成: {filepath}")

            # 3. 同时返回消息内容和文件下载链接
            resp_content = rag_answer
            if file_url:
                fmt_name = "Word文档(.docx)" if fmt == 'docx' else "Markdown(.md)"
                resp_content += f"\n\n📄 {fmt_name}已生成！点击下方卡片下载文件："

            return jsonify({
                "content": resp_content,
                "source": "数据手册助手",
                "file_url": file_url,
                "file_path": file_url,
            })

        if image_data: model = "doubao"
        if not model: model = 'deepseek'
        config = get_model_config()[model]
        headers = {"Authorization": f"Bearer {config['key']}", "Content-Type": "application/json"}

        try:
            if model == "doubao":
                content = []
                if image_data: content.append({"type": "input_image", "image_url": f"data:image/png;base64,{image_data}"})
                content.append({"type": "input_text", "text": question if question else "描述这张图片"})
                payload = {"model": config["model"], "input": [{"role": "user", "content": content}]}
            else:
                payload = {"model": config["model"], "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": question}]}

            resp = requests.post(config["url"], headers=headers, json=payload, timeout=90)
            result = resp.json()
            answer = ""

            if "output" in result:
                for out in result["output"]:
                    if out.get("type") == "message":
                        for c in out.get("content", []):
                            if c.get("type") == "output_text":
                                answer = c.get("text", "").strip()
                                break
                        break
            elif "choices" in result and len(result["choices"]) > 0:
                answer = result["choices"][0]["message"].get("content", "").strip()

            if not answer:
                answer = f"提示：{result.get('error', {}).get('message', '服务暂时无法回复')}"

            return jsonify({"content": answer, "source": {"deepseek": "DeepSeek", "kimi": "Kimi", "doubao": "豆包"}.get(model, "AI")})

        except Exception as e:
            return jsonify({"content": f"服务异常，请稍后再试: {str(e)}", "source": "系统"})

# 🚨 新增：让手机能够下载文件的专属接口
@app.route('/api/download/<path:filename>', methods=['GET'])
def download_file(filename):
    # 按顺序在多个输出目录中查找文件
    for folder in [OUTPUT_FOLDER, 'paper_outputs', 'datasheet_outputs']:
        path = os.path.join(folder, filename)
        if os.path.exists(path):
            return send_from_directory(folder, filename, as_attachment=True)
    return jsonify({"error": "文件不存在"}), 404

# ============ 数据手册 RAG 接口 ============
DATASHEET_FOLDER = 'datasheets'

@app.route('/api/datasheet/upload', methods=['POST'])
def datasheet_upload():
    """上传并解析数据手册PDF"""
    try:
        if 'document' not in request.files:
            return jsonify({"content": "未接收到文件", "source": "错误"}), 400

        file = request.files['document']
        if not file.filename.endswith('.pdf'):
            return jsonify({"content": "只支持PDF文件", "source": "错误"}), 400

        # 保存文件
        safe_filename = f"upload_{int(time.time())}_{file.filename}"
        filepath = os.path.join(DATASHEET_FOLDER, safe_filename)
        os.makedirs(DATASHEET_FOLDER, exist_ok=True)
        file.save(filepath)

        # 加载并索引文档
        result = datasheet_rag.load_document(filepath)

        return jsonify({
            "content": f"数据手册解析完成！\n\n文件名：{result['filename']}\n切分段落：{result['chunks']}块\n芯片型号：{result['metadata'].get('chip_name', '未知')}",
            "source": "数据手册助手",
            "filename": result['filename'],
            "chunks": result['chunks']
        })

    except Exception as e:
        return jsonify({"content": f"处理失败：{str(e)}", "source": "错误"}), 500

@app.route('/api/datasheet/ask', methods=['POST'])
def datasheet_ask():
    """问答接口"""
    try:
        data = request.json
        question = data.get('question', '').strip()

        if not question:
            return jsonify({"content": "请输入问题", "source": "错误"}), 400

        # 调用RAG引擎回答
        result = datasheet_rag.ask(question)

        # 自动检测导出意图
        export_keywords = ['导出', '输出', '生成文档', '生成word', '生成md', '下载']
        need_export = any(kw in question.lower() for kw in export_keywords)

        if need_export and result.get('content'):
            # 自动生成文档并返回下载链接
            fmt = 'md' if 'md' in question.lower() or 'markdown' in question.lower() else 'docx'
            doc_result = datasheet_rag.export_to_document(question, result['content'], fmt)
            server_ip = os.getenv("SERVER_IP", "192.168.2.6")
            server_port = os.getenv("FLASK_PORT", "5000")
            filepath = doc_result.get("filepath", "")
            if filepath and os.path.exists(filepath):
                filename = os.path.basename(filepath)
                doc_result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
                result["file_url"] = doc_result["file_url"]
                result["content"] = f"{result['content']}\n\n📄 文档已自动生成！点击下方链接下载：\n{doc_result['file_url']}"

        return jsonify(result)

    except Exception as e:
        return jsonify({"content": f"处理失败：{str(e)}", "source": "错误"}), 500

@app.route('/api/datasheet/status', methods=['GET'])
def datasheet_status():
    """获取当前加载的文档状态"""
    return jsonify({
        "has_document": datasheet_rag.current_file is not None,
        "filename": datasheet_rag.current_file,
        "chunks": len(datasheet_rag.documents)
    })

if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

    host = os.getenv("FLASK_HOST", "0.0.0.0")
    port = int(os.getenv("FLASK_PORT", 5000))
    server_ip = os.getenv("SERVER_IP", "192.168.2.6")

    print("\n" + "=" * 60)
    print("Backend service started successfully!")
    print("=" * 60)
    print(f"Listen address: {host}:{port}")
    print(f"Flutter access URL: http://{server_ip}:{port}/api/chat")
    print("=" * 60 + "\n")

    app.run(host=host, port=port, debug=False)


# ============ 论文助手接口 ============
PAPER_FOLDER = 'paper_uploads'
os.makedirs(PAPER_FOLDER, exist_ok=True)

@app.route('/api/paper/outline', methods=['POST'])
def paper_outline():
    """生成论文大纲"""
    try:
        data = request.json
        topic = data.get('topic', '').strip()
        discipline = data.get('discipline', 'general')

        if not topic:
            return jsonify({"content": "请输入论文主题", "source": "论文助手"})

        result = generate_outline(topic, discipline)

        # 自动检测导出意图
        export_keywords = ['导出', '输出', '生成文档', '生成word', '生成md', '下载']
        need_export = any(kw in topic.lower() for kw in export_keywords)

        if need_export:
            # 自动生成文档并返回下载链接
            fmt = 'md' if 'md' in topic.lower() or 'markdown' in topic.lower() else 'docx'
            doc_result = generate_outline_document(topic, discipline, fmt)
            server_ip = os.getenv("SERVER_IP", "192.168.2.6")
            server_port = os.getenv("FLASK_PORT", "5000")
            filepath = doc_result.get("filepath", "")
            if filepath and os.path.exists(filepath):
                filename = os.path.basename(filepath)
                doc_result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
                result["file_url"] = doc_result["file_url"]
                result["content"] = f"{result['content']}\n\n📄 文档已自动生成！点击下方链接下载：\n{doc_result['file_url']}"

        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"生成失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/polish', methods=['POST'])
def paper_polish():
    """润色文本"""
    try:
        data = request.json
        text = data.get('text', '').strip()

        if not text:
            return jsonify({"content": "请输入要润色的文本", "source": "论文助手"})

        result = polish_text(text)
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"润色失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/rewrite', methods=['POST'])
def paper_rewrite():
    """降重改写"""
    try:
        data = request.json
        text = data.get('text', '').strip()

        if not text:
            return jsonify({"content": "请输入要改写的文本", "source": "论文助手"})

        result = rewrite_text(text)
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"改写失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/references', methods=['POST'])
def paper_references():
    """格式化参考文献"""
    try:
        data = request.json
        text = data.get('text', '').strip()
        format_type = data.get('format', 'gb')

        if not text:
            return jsonify({"content": "请输入文献信息", "source": "论文助手"})

        result = format_references(text, format_type)
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"格式化失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/check-format', methods=['POST'])
def paper_check_format():
    """检查论文格式"""
    try:
        if 'document' not in request.files:
            # 无文档时返回检查清单
            result = check_format()
            return jsonify(result)

        file = request.files['document']
        # 保存文件
        safe_filename = f"upload_{int(time.time())}_{file.filename}"
        filepath = os.path.join(PAPER_FOLDER, safe_filename)
        file.save(filepath)

        result = check_format(filepath)
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"检查失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/template', methods=['GET'])
def paper_template():
    """获取写作模板"""
    try:
        template_type = request.args.get('type', 'outline')
        result = get_template(template_type)
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"获取模板失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/export-outline', methods=['POST'])
def paper_export_outline():
    """导出大纲为文档"""
    try:
        data = request.json
        topic = data.get('topic', '').strip()
        discipline = data.get('discipline', 'general')
        fmt = data.get('format', 'docx')

        if not topic:
            return jsonify({"content": "请输入论文主题", "source": "论文助手"})

        result = generate_outline_document(topic, discipline, fmt)
        # 返回文件下载地址
        server_ip = os.getenv("SERVER_IP", "192.168.2.6")
        server_port = os.getenv("FLASK_PORT", "5000")
        filepath = result.get("filepath", "")
        if filepath and os.path.exists(filepath):
            filename = os.path.basename(filepath)
            result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"导出失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/export-polish', methods=['POST'])
def paper_export_polish():
    """导出润色结果为文档"""
    try:
        data = request.json
        original = data.get('original', '').strip()
        polished = data.get('polished', '').strip()
        fmt = data.get('format', 'docx')

        if not original or not polished:
            return jsonify({"content": "请先生成润色结果", "source": "论文助手"})

        result = generate_polish_document(original, polished, fmt)
        server_ip = os.getenv("SERVER_IP", "192.168.2.6")
        server_port = os.getenv("FLASK_PORT", "5000")
        filepath = result.get("filepath", "")
        if filepath and os.path.exists(filepath):
            filename = os.path.basename(filepath)
            result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"导出失败：{str(e)}", "source": "论文助手"})


@app.route('/api/paper/export-references', methods=['POST'])
def paper_export_references():
    """导出参考文献为文档"""
    try:
        data = request.json
        refs = data.get('refs', '').strip()
        fmt_type = data.get('format', 'gb')

        if not refs:
            return jsonify({"content": "请先生成参考文献", "source": "论文助手"})

        result = generate_references_document(refs, fmt_type)
        server_ip = os.getenv("SERVER_IP", "192.168.2.6")
        server_port = os.getenv("FLASK_PORT", "5000")
        filepath = result.get("filepath", "")
        if filepath and os.path.exists(filepath):
            filename = os.path.basename(filepath)
            result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"导出失败：{str(e)}", "source": "论文助手"})


# ============ 数据手册导出接口 ============
@app.route('/api/datasheet/export-qa', methods=['POST'])
def datasheet_export_qa():
    """导出当前问答为文档"""
    try:
        data = request.json
        question = data.get('question', '').strip()
        answer = data.get('answer', '').strip()
        fmt = data.get('format', 'docx')

        if not question or not answer:
            return jsonify({"content": "没有可导出的问答内容", "source": "错误"})

        result = datasheet_rag.export_to_document(question, answer, fmt)
        server_ip = os.getenv("SERVER_IP", "192.168.2.6")
        server_port = os.getenv("FLASK_PORT", "5000")
        filepath = result.get("filepath", "")
        if filepath and os.path.exists(filepath):
            filename = os.path.basename(filepath)
            result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"导出失败：{str(e)}", "source": "错误"})


@app.route('/api/datasheet/export-full', methods=['POST'])
def datasheet_export_full():
    """导出完整数据手册为文档"""
    try:
        fmt = request.args.get('format', 'docx')
        result = datasheet_rag.export_full_specs(fmt)
        server_ip = os.getenv("SERVER_IP", "192.168.2.6")
        server_port = os.getenv("FLASK_PORT", "5000")
        filepath = result.get("filepath", "")
        if filepath and os.path.exists(filepath):
            filename = os.path.basename(filepath)
            result["file_url"] = f"http://{server_ip}:{server_port}/api/download/{urllib.parse.quote(filename)}"
        return jsonify(result)
    except Exception as e:
        return jsonify({"content": f"导出失败：{str(e)}", "source": "错误"})
